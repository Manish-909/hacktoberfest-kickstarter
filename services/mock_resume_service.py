import json
import io
import PyPDF2
import pdfplumber
from docx import Document
from typing import Dict, Optional, BinaryIO
import re
import time

class MockResumeService:
    """Mock Resume Service for demonstration that simulates AI analysis"""
    
    def __init__(self, api_key: str = None):
        # Predefined lists for skills and interests
        self.SKILLS_LIST = [
            'JavaScript', 'Python', 'Java', 'TypeScript', 'React', 'Vue.js', 'Angular',
            'Node.js', 'Express', 'Django', 'Flask', 'Spring Boot', 'PHP', 'Laravel',
            'Ruby', 'Rails', 'Go', 'Rust', 'C++', 'C#', '.NET', 'Swift', 'Kotlin',
            'HTML', 'CSS', 'SCSS', 'Tailwind', 'Bootstrap', 'jQuery', 'Redux',
            'GraphQL', 'REST API', 'MongoDB', 'PostgreSQL', 'MySQL', 'SQLite',
            'Docker', 'Kubernetes', 'AWS', 'Azure', 'GCP', 'Git', 'Linux', 'Bash'
        ]
        
        self.INTERESTS_LIST = [
            'frontend', 'backend', 'fullstack', 'mobile', 'web', 'api', 'database',
            'devops', 'cloud', 'ai', 'machine-learning', 'data-science', 'blockchain',
            'security', 'testing', 'documentation', 'ui-ux', 'performance', 'accessibility',
            'open-source', 'beginner-friendly', 'hacktoberfest', 'help-wanted', 'bug',
            'feature', 'enhancement', 'refactor', 'optimization', 'tutorial'
        ]
    
    def extract_text_from_pdf(self, pdf_file: BinaryIO) -> str:
        """Extract text from PDF file using pdfplumber for better accuracy"""
        try:
            text_content = []
            
            with pdfplumber.open(pdf_file) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_content.append(page_text)
            
            return "\n".join(text_content)
            
        except Exception as e:
            # Fallback to PyPDF2
            try:
                pdf_file.seek(0)  # Reset file pointer
                pdf_reader = PyPDF2.PdfReader(pdf_file)
                text_content = []
                
                for page in pdf_reader.pages:
                    text_content.append(page.extract_text())
                
                return "\n".join(text_content)
            except Exception as fallback_e:
                raise Exception(f"Failed to extract PDF text: {str(e)}, Fallback error: {str(fallback_e)}")
    
    def extract_text_from_docx(self, docx_file: BinaryIO) -> str:
        """Extract text from DOCX file"""
        try:
            doc = Document(docx_file)
            text_content = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text)
                    if row_text:
                        text_content.append(" | ".join(row_text))
            
            return "\n".join(text_content)
            
        except Exception as e:
            raise Exception(f"Failed to extract DOCX text: {str(e)}")
    
    def extract_text_from_file(self, uploaded_file) -> str:
        """Extract text from uploaded file based on its type"""
        file_type = uploaded_file.type
        
        if file_type == "application/pdf":
            return self.extract_text_from_pdf(uploaded_file)
        elif file_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            return self.extract_text_from_docx(uploaded_file)
        else:
            raise Exception(f"Unsupported file type: {file_type}. Please upload a PDF or DOCX file.")
    
    def analyze_resume_with_ai(self, resume_text: str) -> Dict:
        """Mock AI analysis that simulates intelligent resume parsing"""
        
        # Simulate processing time
        time.sleep(2)
        
        text_lower = resume_text.lower()
        
        # Determine experience level based on keywords and content length
        experience_level = "intermediate"  # Default
        
        # Look for experience indicators
        if any(word in text_lower for word in ["senior", "lead", "principal", "architect", "manager", "5+ years", "6+ years", "7+ years"]):
            experience_level = "advanced"
        elif any(word in text_lower for word in ["junior", "entry", "intern", "graduate", "1 year", "2 years", "new grad"]):
            experience_level = "beginner"
        elif any(word in text_lower for word in ["3 years", "4 years", "5 years", "mid-level", "intermediate"]):
            experience_level = "intermediate"
        
        # Extract skills by looking for them in the resume text
        found_skills = []
        for skill in self.SKILLS_LIST:
            if skill.lower() in text_lower:
                found_skills.append(skill)
        
        # Extract interests based on project descriptions and roles
        found_interests = []
        
        # Map certain keywords to interests
        interest_mapping = {
            'web': ['web', 'frontend', 'backend'],
            'api': ['api', 'backend'],
            'mobile': ['mobile'],
            'react': ['frontend', 'web'],
            'node': ['backend', 'web'],
            'python': ['backend', 'ai', 'data-science'],
            'machine learning': ['ai', 'machine-learning', 'data-science'],
            'docker': ['devops', 'cloud'],
            'aws': ['devops', 'cloud'],
            'database': ['database', 'backend'],
            'frontend': ['frontend', 'web', 'ui-ux'],
            'backend': ['backend', 'api'],
            'fullstack': ['fullstack', 'web', 'frontend', 'backend'],
            'security': ['security'],
            'test': ['testing'],
            'documentation': ['documentation'],
            'open source': ['open-source'],
        }
        
        for keyword, interests in interest_mapping.items():
            if keyword in text_lower:
                found_interests.extend(interests)
        
        # Remove duplicates and ensure they're in our predefined list
        found_interests = list(set([interest for interest in found_interests if interest in self.INTERESTS_LIST]))
        
        # Add some default interests if none found
        if not found_interests:
            found_interests = ['beginner-friendly', 'help-wanted', 'bug', 'enhancement']
        
        # Ensure we have some skills if none found
        if not found_skills:
            # Infer some common skills based on experience level
            if experience_level == "beginner":
                found_skills = ['HTML', 'CSS', 'JavaScript', 'Git']
            elif experience_level == "intermediate":
                found_skills = ['JavaScript', 'Python', 'React', 'Node.js', 'Git', 'HTML', 'CSS']
            else:
                found_skills = ['JavaScript', 'Python', 'React', 'Node.js', 'Docker', 'AWS', 'Git', 'MongoDB']
        
        return {
            "experienceLevel": experience_level,
            "programmingSkills": found_skills[:15],  # Limit to avoid too many selections
            "areasOfInterest": found_interests[:12]   # Limit to avoid too many selections
        }
    
    def process_resume(self, uploaded_file) -> Dict:
        """Complete resume processing pipeline"""
        try:
            # Extract text from file
            resume_text = self.extract_text_from_file(uploaded_file)
            
            if not resume_text.strip():
                raise Exception("No text could be extracted from the resume file")
            
            # Analyze with Mock AI
            profile_data = self.analyze_resume_with_ai(resume_text)
            
            return {
                "success": True,
                "profile": profile_data,
                "resume_text_preview": resume_text[:500] + "..." if len(resume_text) > 500 else resume_text
            }
            
        except Exception as e:
            return {
                "success": False,
                "error": str(e),
                "profile": None
            }
    
    def validate_extracted_profile(self, profile_data: Dict) -> Dict:
        """Validate and clean the extracted profile data"""
        cleaned_profile = {
            "experience_level": "intermediate",  # Default
            "skills": [],
            "interests": []
        }
        
        if profile_data:
            # Validate experience level
            experience = profile_data.get("experienceLevel", "").lower()
            if experience in ["beginner", "intermediate", "advanced"]:
                cleaned_profile["experience_level"] = experience
            
            # Validate skills (only include those in our predefined list)
            skills = profile_data.get("programmingSkills", [])
            if isinstance(skills, list):
                cleaned_profile["skills"] = [
                    skill for skill in skills if skill in self.SKILLS_LIST
                ]
            
            # Validate interests (only include those in our predefined list)
            interests = profile_data.get("areasOfInterest", [])
            if isinstance(interests, list):
                cleaned_profile["interests"] = [
                    interest for interest in interests if interest in self.INTERESTS_LIST
                ]
        
        return cleaned_profile
