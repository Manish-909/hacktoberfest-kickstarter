import json
import io
import PyPDF2
import pdfplumber
from docx import Document
from typing import Dict, Optional, BinaryIO
import streamlit as st
import google.generativeai as genai

class ResumeService:
    """Service for processing resume files and extracting profile information using Gemini AI"""
    
    def __init__(self, api_key: str = "AIzaSyCZ7sFBZKmuj4Tk4Md7NF-G5W4pI0Hxz7A"):
        genai.configure(api_key=api_key)
        self.model = "models/text-bison-001"
        
        # Predefined lists for skills and interests
        self.SKILLS_LIST = [
            'JavaScript', 'Python', 'Java', 'TypeScript', 'React', 'Vue.js', 'Angular',
            'Node.js', 'Express', 'Django', 'Flask', 'Spring Boot', 'PHP', 'Laravel',
            'Ruby', 'Rails', 'Go', 'Rust', 'C++', 'C#', '.NET', 'Swift', 'Kotlin',
            'HTML', 'CSS', 'SCSS', 'Tailwind', 'Bootstrap', 'jQuery', 'Redux',
            'GraphQL', 'REST API', 'MongoDB', 'PostgreSQL', 'MySQL', 'SQLite',
            'Docker', 'Kubernetes', 'AWS', 'Azure', 'GCP', 'Git', 'Linux', 'Bash'
        ]
        
        self.INTERESTS_LIST = [
            'frontend', 'backend', 'fullstack', 'mobile', 'web', 'api', 'database',
            'devops', 'cloud', 'ai', 'machine-learning', 'data-science', 'blockchain',
            'security', 'testing', 'documentation', 'ui-ux', 'performance', 'accessibility',
            'open-source', 'beginner-friendly', 'hacktoberfest', 'help-wanted', 'bug',
            'feature', 'enhancement', 'refactor', 'optimization', 'tutorial'
        ]
    
    def extract_text_from_pdf(self, pdf_file: BinaryIO) -> str:
        """Extract text from PDF file using pdfplumber for better accuracy"""
        try:
            text_content = []
            
            with pdfplumber.open(pdf_file) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_content.append(page_text)
            
            return "\n".join(text_content)
            
        except Exception as e:
            # Fallback to PyPDF2
            try:
                pdf_file.seek(0)  # Reset file pointer
                pdf_reader = PyPDF2.PdfReader(pdf_file)
                text_content = []
                
                for page in pdf_reader.pages:
                    text_content.append(page.extract_text())
                
                return "\n".join(text_content)
            except Exception as fallback_e:
                raise Exception(f"Failed to extract PDF text: {str(e)}, Fallback error: {str(fallback_e)}")
    
    def extract_text_from_docx(self, docx_file: BinaryIO) -> str:
        """Extract text from DOCX file"""
        try:
            doc = Document(docx_file)
            text_content = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text)
                    if row_text:
                        text_content.append(" | ".join(row_text))
            
            return "\n".join(text_content)
            
        except Exception as e:
            raise Exception(f"Failed to extract DOCX text: {str(e)}")
    
    def extract_text_from_file(self, uploaded_file) -> str:
        """Extract text from uploaded file based on its type"""
        file_type = uploaded_file.type
        
        if file_type == "application/pdf":
            return self.extract_text_from_pdf(uploaded_file)
        elif file_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            return self.extract_text_from_docx(uploaded_file)
        else:
            raise Exception(f"Unsupported file type: {file_type}. Please upload a PDF or DOCX file.")
    
    def analyze_resume_with_gemini(self, resume_text: str) -> Dict:
        """Analyze resume text using Gemini AI to extract profile information"""
        
        prompt = f"""Act as an expert career counselor AI and resume interpreter. Your primary function is to analyze the text of a resume to extract key information about a candidate's professional profile.

Analyze the provided resume text to identify the candidate's:
1. Experience Level: Categorize their overall professional experience.
2. Programming Skills: Identify all programming languages, frameworks, and technologies they know.
3. Areas of Interest: Infer their professional interests based on their project descriptions, summary, and experience.

You must strictly use the predefined lists for Programming Skills and Areas of Interest.

SKILLS LIST:
{self.SKILLS_LIST}

INTERESTS LIST:
{self.INTERESTS_LIST}

Steps to be followed:
1. Thoroughly read and parse the entire resume text provided.
2. Based on job titles (e.g., "Senior Developer", "Intern"), years of experience mentioned, and the scope of responsibilities, determine the Experience Level. Classify it as one of the following: beginner, intermediate, or advanced.
3. Scan the text specifically for technical skills. Match the skills you find against the provided SKILLS LIST. Create a list of all matching skills.
4. Analyze the project descriptions, professional summary, and job roles to deduce the candidate's professional interests. Match these inferred interests against the INTERESTS LIST. Create a list of all matching interests.

Format the final output as a single JSON object.

Your output must be a single, clean JSON object. Do not include any introductory text, explanations, or apologies. The JSON object should have three keys: experienceLevel, programmingSkills, and areasOfInterest.

Resume Text:
{resume_text}

Expected JSON Output Format:
{{
"experienceLevel": "beginner|intermediate|advanced",
"programmingSkills": ["skill1", "skill2", ...],
"areasOfInterest": ["interest1", "interest2", ...]
}}"""

        try:
            response = genai.generate_text(
                model=self.model,
                prompt=prompt,
                temperature=0.3,
                max_output_tokens=1000,
            )
            
            # Extract the response text
            response_text = response.result.strip() if response.result else ""
            
            # Try to parse the JSON response
            try:
                # Remove any markdown code block formatting
                if response_text.startswith("```"):
                    response_text = response_text.split("```")[1]
                if response_text.startswith("json"):
                    response_text = response_text[4:].strip()
                
                profile_data = json.loads(response_text)
                
                # Validate the response structure
                if not all(key in profile_data for key in ["experienceLevel", "programmingSkills", "areasOfInterest"]):
                    raise ValueError("Invalid response structure from Gemini")
                
                # Ensure experience level is lowercase and valid
                valid_experience_levels = ["beginner", "intermediate", "advanced"]
                if profile_data["experienceLevel"].lower() not in valid_experience_levels:
                    profile_data["experienceLevel"] = "intermediate"  # Default fallback
                else:
                    profile_data["experienceLevel"] = profile_data["experienceLevel"].lower()
                
                # Validate skills and interests are lists
                if not isinstance(profile_data["programmingSkills"], list):
                    profile_data["programmingSkills"] = []
                if not isinstance(profile_data["areasOfInterest"], list):
                    profile_data["areasOfInterest"] = []
                
                return profile_data
                
            except json.JSONDecodeError as e:
                raise Exception(f"Failed to parse Gemini response as JSON: {str(e)}")
                
        except Exception as e:
            raise Exception(f"Gemini API error: {str(e)}")
    
    def process_resume(self, uploaded_file) -> Dict:
        """Complete resume processing pipeline"""
        try:
            # Extract text from file
            resume_text = self.extract_text_from_file(uploaded_file)
            
            if not resume_text.strip():
                raise Exception("No text could be extracted from the resume file")
            
            # Analyze with Gemini
            profile_data = self.analyze_resume_with_gemini(resume_text)
            
            return {
                "success": True,
                "profile": profile_data,
                "resume_text_preview": resume_text[:500] + "..." if len(resume_text) > 500 else resume_text
            }
            
        except Exception as e:
            return {
                "success": False,
                "error": str(e),
                "profile": None
            }
    
    def validate_extracted_profile(self, profile_data: Dict) -> Dict:
        """Validate and clean the extracted profile data"""
        cleaned_profile = {
            "experienceLevel": "intermediate",  # Default
            "programmingSkills": [],
            "areasOfInterest": []
        }
        
        if profile_data:
            # Validate experience level
            experience = profile_data.get("experienceLevel", "").lower()
            if experience in ["beginner", "intermediate", "advanced"]:
                cleaned_profile["experienceLevel"] = experience
            
            # Validate skills (only include those in our predefined list)
            skills = profile_data.get("programmingSkills", [])
            if isinstance(skills, list):
                cleaned_profile["programmingSkills"] = [
                    skill for skill in skills if skill in self.SKILLS_LIST
                ]
            
            # Validate interests (only include those in our predefined list)
            interests = profile_data.get("areasOfInterest", [])
            if isinstance(interests, list):
                cleaned_profile["areasOfInterest"] = [
                    interest for interest in interests if interest in self.INTERESTS_LIST
                ]
        
        return cleaned_profile
